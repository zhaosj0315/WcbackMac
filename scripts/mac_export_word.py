#!/usr/bin/env python3
"""
Mac 版本 Word 文档导出器
"""
import sqlite3
from pathlib import Path
from datetime import datetime
from typing import Optional, List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx 未安装，请运行: pip install python-docx")

from mac_message_utils import MacMediaResolver, parse_message


class MacWordExporter:
    def __init__(self, db_dir: str = "app/Database/MacMsg", contact_mapper=None):
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        
        self.db_dir = Path(db_dir)
        self.contact_mapper = contact_mapper
        self.message_dir = self.db_dir / "message"
        self.media_resolver = MacMediaResolver(self.db_dir)
    
    def export_conversation(self, table_name: str, output_path: str, 
                           limit: Optional[int] = None):
        """导出单个会话到 Word"""
        db_files = self._find_message_dbs(table_name)
        if not db_files:
            print(f"❌ 未找到表 {table_name}")
            return

        messages = []
        for db_file in db_files:
            conn = sqlite3.connect(str(db_file))
            cursor = conn.cursor()
            cursor.execute(f"""
                SELECT local_id, real_sender_id, local_type, create_time, message_content, sort_seq
                FROM {table_name}
            """)
            messages.extend(cursor.fetchall())
            conn.close()
        messages.sort(key=lambda row: (row[3] or 0, row[0] or 0))
        if limit:
            messages = messages[:limit]
        
        self._generate_docx(messages, table_name, output_path)
        print(f"✅ 导出 {len(messages)} 条消息到 {output_path}")
    
    def export_all_conversations(self, output_dir: str = "data/word_export",
                                limit_per_chat: Optional[int] = None):
        """导出所有会话"""
        output = Path(output_dir)
        output.mkdir(parents=True, exist_ok=True)
        
        tables = self._get_all_message_tables()
        print(f"📊 找到 {len(tables)} 个会话")
        
        # 构建表名到 wxid 的映射
        import hashlib
        table_to_wxid = {}
        if self.contact_mapper:
            all_wxids = list(self.contact_mapper.wxid_to_remark.keys()) + list(self.contact_mapper.chatroom_to_name.keys())
            for wxid in all_wxids:
                table_hash = hashlib.md5(wxid.encode()).hexdigest()
                table_name = f"Msg_{table_hash}"
                table_to_wxid[table_name] = wxid
        
        for i, (db_file, table_name) in enumerate(tables, 1):
            try:
                # 从表名反查 wxid
                wxid = table_to_wxid.get(table_name, table_name.replace('Msg_', ''))
                
                if self.contact_mapper:
                    if '@chatroom' in wxid:
                        display_name = self.contact_mapper.get_chatroom_name(wxid)
                    else:
                        display_name = self.contact_mapper.get_display_name(wxid)
                else:
                    display_name = wxid
                
                # 安全的文件名
                safe_name = "".join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in display_name)
                output_file = output / f"{safe_name}.docx"
                
                # 如果文件名重复，添加 wxid 后缀
                if output_file.exists():
                    output_file = output / f"{safe_name}_{wxid[:8]}.docx"
                
                self.export_conversation(table_name, str(output_file), limit_per_chat)
                
                if i % 10 == 0:
                    print(f"  进度: {i}/{len(tables)}")
            except Exception as e:
                print(f"  ⚠️  导出失败 {table_name}: {e}")
        
        print(f"\n✅ 全部导出完成: {output.absolute()}")
    
    def _find_message_dbs(self, table_name: str) -> List[Path]:
        """查找包含指定表的全部数据库分片"""
        db_files = []
        for db_file in sorted(self.message_dir.glob("message_*.db")):
            try:
                conn = sqlite3.connect(str(db_file))
                cursor = conn.cursor()
                cursor.execute(f"SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
                if cursor.fetchone():
                    db_files.append(db_file)
                conn.close()
            except:
                pass
        return db_files
    
    def _get_all_message_tables(self) -> List[tuple]:
        """获取所有消息表"""
        tables = []
        seen = set()
        for db_file in sorted(self.message_dir.glob("message_*.db")):
            try:
                conn = sqlite3.connect(str(db_file))
                cursor = conn.cursor()
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name LIKE 'Msg_%'")
                for row in cursor.fetchall():
                    if row[0] not in seen:
                        seen.add(row[0])
                        tables.append((db_file, row[0]))
                conn.close()
            except Exception as e:
                print(f"⚠️  读取 {db_file.name} 失败: {e}")
        return tables
    
    def _generate_docx(self, messages: List[tuple], table_name: str, output_path: str):
        """生成 Word 文档"""
        wxid = table_name.replace('Msg_', '')
        display_name = self.contact_mapper.get_display_name(wxid) if self.contact_mapper else wxid
        
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(10.5)
        
        # 标题
        title = doc.add_heading(f'聊天记录 - {display_name}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 元信息
        doc.add_paragraph(f'微信号: {wxid}')
        doc.add_paragraph(f'消息数量: {len(messages)}')
        doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('=' * 60)
        
        # 消息内容
        for msg in messages:
            local_id, real_sender_id, msg_type, create_time, content, sort_seq = msg
            
            time_str = datetime.fromtimestamp(create_time).strftime('%Y-%m-%d %H:%M:%S')
            sender = "我" if real_sender_id == 0 else display_name
            
            # 时间和发送者
            p = doc.add_paragraph()
            run = p.add_run(f'[{time_str}] {sender}')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            
            # 消息内容
            parsed = parse_message(
                msg_type,
                content,
                table_name=table_name,
                local_id=local_id,
                create_time=create_time,
                sort_seq=sort_seq,
                resolver=self.media_resolver,
            )

            p = doc.add_paragraph(parsed.text)
            
            # 发送的消息用不同颜色
            if real_sender_id == 0 and p.runs:
                p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

            media_dir = Path(output_path).with_suffix("").parent / f"{Path(output_path).stem}_media"
            if parsed.media_kind == "image" and parsed.media_path:
                try:
                    doc.add_picture(parsed.media_path, width=Inches(3.2))
                except Exception as exc:
                    doc.add_paragraph(f"[图片文件存在但写入 Word 失败: {exc}] {parsed.media_path}")
            elif parsed.media_kind == "video" and parsed.media_path:
                copied = self.media_resolver.copy_media(parsed.media_path, media_dir, prefix=f"{local_id}_")
                doc.add_paragraph(f"[视频文件] {copied}")
            elif parsed.media_kind == "voice":
                voice_data = self.media_resolver.get_voice_data(table_name, local_id, create_time)
                if voice_data:
                    voice_path = self.media_resolver.write_voice(
                        voice_data, media_dir, f"{local_id}_{create_time}.silk"
                    )
                    doc.add_paragraph(f"[语音文件] {voice_path}")
            
            # 添加间距
            p.paragraph_format.space_after = Pt(6)
        
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
    
def main():
    import argparse
    parser = argparse.ArgumentParser(description='Mac 微信 Word 导出')
    parser.add_argument('--db-dir', default='app/Database/MacMsg', help='数据库目录')
    parser.add_argument('--output', default='data/word_export', help='输出目录')
    parser.add_argument('--mapping', default='data/mac_contact_mapping.json', help='联系人映射文件')
    parser.add_argument('--limit', type=int, help='每个会话最多导出消息数')
    parser.add_argument('--table', help='只导出指定表')
    args = parser.parse_args()
    
    from mac_contact_mapper import MacContactMapper
    mapper = None
    if Path(args.mapping).exists():
        mapper = MacContactMapper.load_mapping(args.mapping)
    
    exporter = MacWordExporter(args.db_dir, mapper)
    
    if args.table:
        output_file = Path(args.output) / f"{args.table}.docx"
        exporter.export_conversation(args.table, str(output_file), args.limit)
    else:
        exporter.export_all_conversations(args.output, args.limit)


if __name__ == '__main__':
    main()
