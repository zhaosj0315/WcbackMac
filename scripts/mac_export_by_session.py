#!/usr/bin/env python3
"""
Mac 版本完整导出，完全复刻 Windows 版本结构：
  data/
  ├── image/ video/ voice/ emoji/ files/ music/  ← 全量媒体
  ├── 朋友圈/
  ├── 聊天统计/
  └── 聊天记录/
      └── 昵称(wxid)/
          ├── 昵称_chat.txt
          ├── 昵称.html
          ├── 昵称.csv
          ├── 昵称.txt
          ├── 昵称_N.docx
          ├── 昵称_train.json / 昵称_dev.json
          ├── avatar/
          ├── image/ voice/ video/ emoji/ file/ music/ icon/
"""
import argparse
import csv
import hashlib
import json
import shutil
import sqlite3
import sys
from datetime import datetime
from io import StringIO
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

try:
    import zstd
except ImportError:
    zstd = None

from mac_message_utils import MacMediaResolver, parse_message, silk_to_wav
from mac_export_media import detect_ext


# ── 工具函数 ──────────────────────────────────────────────────────────────────

def _load_contact_map(mapping_path: Path) -> dict[str, str]:
    if not mapping_path.exists():
        return {}
    raw = json.loads(mapping_path.read_text(encoding="utf-8"))
    result = {}
    if "contacts" in raw:
        result.update(raw["contacts"])
    if "chatrooms" in raw:
        result.update(raw["chatrooms"])
    if not result:
        result = raw
    return result


def _find_message_dbs(msg_dir: Path, table: str) -> list[Path]:
    result = []
    for db in sorted(msg_dir.glob("message_*.db")):
        if "fts" in db.name:
            continue
        try:
            conn = sqlite3.connect(db)
            if conn.execute(
                "select name from sqlite_master where type='table' and name=?", (table,)
            ).fetchone():
                result.append(db)
            conn.close()
        except sqlite3.Error:
            continue
    return result


def _load_rows(msg_dir: Path, table: str, limit: int = 0) -> list[tuple]:
    rows = []
    for db in _find_message_dbs(msg_dir, table):
        conn = sqlite3.connect(db)
        cur = conn.cursor()
        cur.execute(
            f'select local_id, local_type, create_time, real_sender_id, '
            f'message_content, sort_seq from "{table}" where create_time > 0'
        )
        rows.extend(cur.fetchall())
        conn.close()
    rows.sort(key=lambda r: (r[2], r[0]))
    return rows[:limit] if limit else rows


def _safe_dir_name(name: str) -> str:
    for ch in r'\/:*?"<>|':
        name = name.replace(ch, "_")
    return name


def _fetch_avatar(wxid: str, db_dir: Path) -> bytes | None:
    """从 contact.db 获取头像 URL，尝试下载"""
    contact_db = db_dir.parent / "contact" / "contact.db"
    if not contact_db.exists():
        return None
    try:
        conn = sqlite3.connect(contact_db)
        row = conn.execute(
            "select small_head_url from contact where username=? limit 1", (wxid,)
        ).fetchone()
        conn.close()
        if row and row[0]:
            import urllib.request
            with urllib.request.urlopen(row[0], timeout=5) as resp:
                return resp.read()
    except Exception:
        pass
    return None


# ── 单会话导出 ────────────────────────────────────────────────────────────────

def export_session(
    table: str,
    msg_dir: Path,
    out_dir: Path,
    display_name: str,
    wxid: str,
    my_wxid: str,
    resolver: MacMediaResolver,
    my_name: str = "我",
    limit: int = 0,
) -> int:
    rows = _load_rows(msg_dir, table, limit)
    if not rows:
        return 0

    out_dir.mkdir(parents=True, exist_ok=True)
    for sub in ("image", "voice", "video", "emoji", "file", "music", "icon", "avatar"):
        (out_dir / sub).mkdir(exist_ok=True)

    # 头像
    for uid in (wxid, my_wxid):
        if uid:
            avatar_path = out_dir / "avatar" / f"{uid}.png"
            if not avatar_path.exists():
                data = _fetch_avatar(uid, msg_dir)
                if data:
                    avatar_path.write_bytes(data)

    # 解析所有消息
    parsed_rows = []
    for local_id, msg_type, create_time, sender_id, content, sort_seq in rows:
        dt = datetime.fromtimestamp(create_time)

        # 发送者识别：用分片 Name2Id（每个分片 rowid 独立）
        # _load_rows 没有记录来源分片，用 resolver 的 shard_map 按 my_wxid 判断
        is_me = sender_id in resolver._my_rowids
        sender_name = my_name if is_me else display_name

        parsed = parse_message(
            msg_type, content,
            table_name=table, local_id=local_id,
            create_time=create_time, sort_seq=sort_seq,
            resolver=resolver,
        )
        # 复制媒体文件到会话子目录
        media_label = parsed.text
        voice_fname = None
        if parsed.media_kind == "image":
            # 用修复后的 find_image（含 create_time 降级匹配）
            src_path = resolver.find_image_with_fallback(table, local_id, create_time, sort_seq)
            if src_path and Path(src_path).exists():
                src = Path(src_path)
                ext = detect_ext(src.read_bytes()) if src.suffix.lower() == ".dat" else src.suffix
                dest_name = f"{local_id}_{create_time}{ext or src.suffix}"
                dest = out_dir / "image" / dest_name
                if not dest.exists():
                    if src.suffix.lower() == ".dat":
                        dest.write_bytes(src.read_bytes())
                    else:
                        shutil.copy2(src, dest)
                media_label = f"[图片:{dest_name}]"
            else:
                media_label = "[图片]"
        elif parsed.media_kind == "voice":
            voice_data = resolver.get_voice_data(table, local_id, create_time)
            if voice_data:
                wav = silk_to_wav(voice_data)
                voice_fname = f"{local_id}_{create_time}.wav" if wav else f"{local_id}_{create_time}.silk"
                dest = out_dir / "voice" / voice_fname
                dest.write_bytes(wav if wav else voice_data)
                dur = f" {parsed.voice_length_ms/1000:.1f}s" if parsed.voice_length_ms else ""
                media_label = f"[语音{dur}:{voice_fname}]"
        elif parsed.media_kind == "video":
            # 用修复后的 find_video（含 create_time 降级匹配）
            src_path = resolver.find_video(table, local_id, create_time, sort_seq)
            if src_path and Path(src_path).exists():
                src = Path(src_path)
                dest = out_dir / "video" / src.name
                if not dest.exists():
                    shutil.copy2(src, dest)
                media_label = f"[视频:{src.name}]"
            else:
                media_label = "[视频]"
        elif parsed.media_kind == "emoji":
            media_label = "[表情包]"

        parsed_rows.append({
            "local_id": local_id,
            "msg_type": msg_type,
            "create_time": create_time,
            "dt": dt,
            "sender_id": sender_id,
            "is_me": is_me,
            "sender_name": sender_name,
            "sender_wxid": my_wxid if is_me else wxid,
            "text": parsed.text,
            "media_label": media_label,
            "voice_fname": voice_fname,
            "parsed": parsed,
        })

    _write_chat_txt(out_dir, display_name, parsed_rows)
    _write_csv(out_dir, display_name, wxid, parsed_rows)
    _write_txt(out_dir, display_name, parsed_rows)
    _write_html(out_dir, display_name, wxid, parsed_rows)
    _write_docx(out_dir, display_name, parsed_rows)
    _write_ai_json(out_dir, display_name, parsed_rows, my_name)
    return len(rows)


# ── 各格式写入 ────────────────────────────────────────────────────────────────

def _write_chat_txt(out_dir: Path, name: str, rows: list[dict]) -> None:
    """Windows 格式 _chat.txt：**日期** 分隔，昵称:内容"""
    lines = []
    last_date = None
    for r in rows:
        date_str = r["dt"].strftime("%Y-%m-%d")
        if date_str != last_date:
            lines.append(f"\n********************{date_str}********************\n")
            last_date = date_str
        lines.append(f"{r['sender_name']}:{r['media_label']} \n")
    (out_dir / f"{name}_chat.txt").write_text("\n".join(lines), encoding="utf-8")


def _write_csv(out_dir: Path, name: str, wxid: str, rows: list[dict]) -> None:
    """CSV 格式，与 Windows 版本字段一致"""
    buf = StringIO()
    writer = csv.writer(buf)
    writer.writerow(["localId", "Type", "IsSender", "CreateTime", "Status",
                     "StrContent", "StrTime", "Remark", "NickName", "Sender"])
    for r in rows:
        writer.writerow([
            r["local_id"], r["msg_type"],
            1 if r.get("is_me", r["sender_id"] == 0) else 0,
            r["create_time"], 0,
            r["text"],
            r["dt"].strftime("%Y-%m-%d %H:%M:%S"),
            r["sender_name"], r["sender_name"], r["sender_wxid"],
        ])
    (out_dir / f"{name}.csv").write_text(buf.getvalue(), encoding="utf-8-sig")


def _write_txt(out_dir: Path, name: str, rows: list[dict]) -> None:
    """TXT 格式：时间 发送者\n内容"""
    lines = []
    for r in rows:
        lines.append(f"{r['dt'].strftime('%Y-%m-%d %H:%M:%S')} {r['sender_name']}")
        lines.append(r["media_label"])
        lines.append("")
    (out_dir / f"{name}.txt").write_text("\n".join(lines), encoding="utf-8")


def _write_html(out_dir: Path, name: str, wxid: str, rows: list[dict]) -> None:
    """HTML 格式，微信气泡样式"""
    import html as html_mod
    msgs_html = []
    for r in rows:
        is_me = r.get("is_me", r["sender_id"] == 0)
        cls = "sent" if is_me else "recv"
        p = r["parsed"]
        if p.media_kind == "image":
            img_name = Path(r["media_label"].split(":")[-1].rstrip("]")).name if ":" in r["media_label"] else ""
            body = f'<img src="image/{html_mod.escape(img_name)}" style="max-width:200px">' if img_name else html_mod.escape(r["text"])
        elif p.media_kind == "voice" and r.get("voice_fname"):
            body = f'<audio controls src="voice/{html_mod.escape(r["voice_fname"])}"></audio>'
        elif p.media_kind == "video":
            vid_name = Path(r["media_label"].split(":")[-1].rstrip("]")).name if ":" in r["media_label"] else ""
            body = f'<video controls src="video/{html_mod.escape(vid_name)}" style="max-width:240px"></video>' if vid_name else html_mod.escape(r["text"])
        else:
            body = html_mod.escape(r["text"])
        time_str = r["dt"].strftime("%Y-%m-%d %H:%M:%S")
        msgs_html.append(
            f'<div class="msg {cls}"><div class="name">{html_mod.escape(r["sender_name"])}</div>'
            f'<div class="bubble">{body}</div>'
            f'<div class="time">{time_str}</div></div>'
        )
    html_content = f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">
<title>{html_mod.escape(name)}</title>
<style>
body{{font-family:PingFang SC,Arial,sans-serif;background:#f0f2f5;margin:0;padding:16px}}
.msg{{display:flex;flex-direction:column;margin:8px 0;max-width:70%}}
.msg.sent{{align-self:flex-end;align-items:flex-end;margin-left:auto}}
.msg.recv{{align-self:flex-start;align-items:flex-start}}
.name{{font-size:12px;color:#999;margin-bottom:2px}}
.bubble{{padding:8px 12px;border-radius:8px;font-size:14px;word-break:break-word}}
.sent .bubble{{background:#95ec69}}.recv .bubble{{background:#fff}}
.time{{font-size:11px;color:#bbb;margin-top:2px}}
img,video{{border-radius:4px;display:block}}
</style></head><body>
{''.join(msgs_html)}
</body></html>"""
    (out_dir / f"{name}.html").write_text(html_content, encoding="utf-8")


def _write_docx(out_dir: Path, name: str, rows: list[dict]) -> None:
    """Word 格式，每 500 条一个文件"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
    except ImportError:
        return
    CHUNK = 500
    for i, start in enumerate(range(0, len(rows), CHUNK), 1):
        chunk = rows[start: start + CHUNK]
        doc = Document()
        doc.add_heading(name, level=1)
        for r in chunk:
            p = doc.add_paragraph()
            run = p.add_run(f"{r['sender_name']}  {r['dt'].strftime('%Y-%m-%d %H:%M:%S')}\n")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            parsed = r["parsed"]
            if parsed.media_kind == "image" and parsed.media_path:
                try:
                    doc.add_picture(parsed.media_path, width=Inches(2.5))
                except Exception:
                    doc.add_paragraph(r["media_label"])
            else:
                content_run = p.add_run(r["media_label"])
                if r["sender_id"] == 0:
                    content_run.font.color.rgb = RGBColor(0, 0x80, 0)
        doc.save(str(out_dir / f"{name}_{i}.docx"))


def _write_ai_json(out_dir: Path, name: str, rows: list[dict], my_name: str) -> None:
    """AI 训练数据，对话格式"""
    conversations = []
    window = []
    for r in rows:
        if r["msg_type"] not in (1,) or not r["text"].strip() or r["text"].startswith("<"):
            continue
        role = "assistant" if r["sender_id"] == 0 else "user"
        window.append({"role": role, "content": r["text"].strip()})
        if len(window) >= 6:
            conversations.append({"conversations": window})
            window = window[3:]  # 50% 重叠
    if len(window) >= 2:
        conversations.append({"conversations": window})
    split = max(1, int(len(conversations) * 0.9))
    (out_dir / f"{name}_train.json").write_text(
        json.dumps(conversations[:split], ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (out_dir / f"{name}_dev.json").write_text(
        json.dumps(conversations[split:], ensure_ascii=False, indent=2), encoding="utf-8"
    )


# ── 全量媒体目录 ──────────────────────────────────────────────────────────────

def export_global_media(data_root: Path, db_dir: Path) -> None:
    """把 Mac 微信 msg/ 目录的全量媒体复制到 data/ 顶层（与 Windows 结构一致）"""
    xwechat = Path.home() / "Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files"
    msg_dirs = sorted(xwechat.glob("wxid_*/msg"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not msg_dirs:
        return
    msg_dir = msg_dirs[0]

    mapping = {
        "attach": data_root / "image",
        "video": data_root / "video",
        "file": data_root / "files",
    }
    for src_name, dest_dir in mapping.items():
        src = msg_dir / src_name
        if not src.exists():
            continue
        dest_dir.mkdir(parents=True, exist_ok=True)
        count = 0
        for f in src.rglob("*"):
            if not f.is_file():
                continue
            if f.suffix.lower() == ".dat":
                data = f.read_bytes()
                ext = detect_ext(data)
                dest = dest_dir / (f.stem + (ext or ".dat"))
                if not dest.exists():
                    dest.write_bytes(data)
            else:
                dest = dest_dir / f.name
                if not dest.exists():
                    shutil.copy2(f, dest)
            count += 1
        if count:
            print(f"  {src_name}/ → {dest_dir.name}/: {count:,} 个文件")

    # voice：从 media_0.db 的 VoiceInfo 表批量提取
    media_db = db_dir.parent / "message" / "media_0.db"
    if media_db.exists():
        voice_dest = data_root / "voice"
        voice_dest.mkdir(parents=True, exist_ok=True)
        try:
            conn = sqlite3.connect(media_db)
            rows = conn.execute(
                "select chat_name_id, local_id, create_time, voice_buf from VoiceInfo where voice_buf is not null"
            ).fetchall()
            conn.close()
            count = 0
            for chat_name_id, local_id, create_time, voice_buf in rows:
                wav = silk_to_wav(voice_buf)
                fname = f"{chat_name_id}_{local_id}_{create_time}.wav" if wav else f"{chat_name_id}_{local_id}_{create_time}.silk"
                dest = voice_dest / fname
                if not dest.exists():
                    dest.write_bytes(wav if wav else voice_buf)
                    count += 1
            if count:
                print(f"  voice: {count:,} 条语音")
        except Exception as e:
            print(f"⚠️  voice 导出失败: {e}")


# ── 朋友圈和聊天统计 ──────────────────────────────────────────────────────────

def export_sns_dir(data_root: Path, db_dir: Path) -> None:
    sns_dir = data_root / "朋友圈"
    sns_dir.mkdir(parents=True, exist_ok=True)
    sns_db = db_dir.parent / "sns" / "sns.db"
    if sns_db.exists():
        import subprocess
        result = subprocess.run(
            [sys.executable, str(ROOT_DIR / "scripts" / "mac_export_sns.py"),
             "--db", str(sns_db), "--output", str(sns_dir / "sns.json")],
        )
        if result.returncode != 0:
            print(f"⚠️  export_sns_dir 失败，returncode={result.returncode}")


def export_stats_dir(data_root: Path, db_dir: Path) -> None:
    stats_dir = data_root / "聊天统计"
    stats_dir.mkdir(parents=True, exist_ok=True)
    import subprocess
    result = subprocess.run(
        [sys.executable, str(ROOT_DIR / "scripts" / "mac_chat_analysis.py"),
         "--db-dir", str(db_dir),
         "--output", str(stats_dir / "analysis.json")],
    )
    if result.returncode != 0:
        print(f"⚠️  export_stats_dir 失败，returncode={result.returncode}")


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Mac 完整导出（完全复刻 Windows 结构）")
    parser.add_argument("--db-dir", default="app/DataBase/MacMsg/message")
    parser.add_argument("--output", default="data", help="输出根目录（data/）")
    parser.add_argument("--mapping", default="data/mac_contact_mapping.json")
    parser.add_argument("--my-name", default="我")
    parser.add_argument("--my-wxid", default="wxid_nagmkhfzh8ok22")
    parser.add_argument("--limit", type=int, default=0)
    parser.add_argument("--wxid", help="只导出指定 wxid")
    parser.add_argument("--no-global-media", action="store_true", help="跳过全量媒体复制")
    args = parser.parse_args()

    msg_dir = ROOT_DIR / args.db_dir
    data_root = ROOT_DIR / args.output
    resolver = MacMediaResolver(msg_dir.parent)

    contact_map = _load_contact_map(ROOT_DIR / args.mapping)
    print(f"✅ 联系人映射: {len(contact_map)} 个")

    table_to_wxid = {
        f"Msg_{hashlib.md5(wxid.encode()).hexdigest()}": wxid
        for wxid in contact_map
    }

    # 全量媒体
    if not args.no_global_media:
        print("📦 复制全量媒体文件...")
        export_global_media(data_root, msg_dir)

    # 朋友圈 & 统计
    print("📸 导出朋友圈...")
    export_sns_dir(data_root, msg_dir)
    print("📊 导出聊天统计...")
    export_stats_dir(data_root, msg_dir)

    # 收集会话表
    tables: set[str] = set()
    for db in sorted(msg_dir.glob("message_*.db")):
        if "fts" in db.name:
            continue
        try:
            conn = sqlite3.connect(db)
            for (t,) in conn.execute(
                "select name from sqlite_master where type='table' and name like 'Msg_%'"
            ):
                tables.add(t)
            conn.close()
        except sqlite3.Error:
            continue

    if args.wxid:
        target = f"Msg_{hashlib.md5(args.wxid.encode()).hexdigest()}"
        tables = {target} if target in tables else set()

    out_root = data_root / "聊天记录"
    print(f"💬 导出 {len(tables)} 个会话...")
    total_msgs = total_sessions = 0

    for table in sorted(tables):
        wxid = table_to_wxid.get(table, "")
        display_name = contact_map.get(wxid, wxid or table.replace("Msg_", "")[:16])
        dir_name = _safe_dir_name(f"{display_name}({wxid})" if wxid else display_name)
        out_dir = out_root / dir_name

        count = export_session(
            table, msg_dir, out_dir, display_name, wxid,
            args.my_wxid, resolver, args.my_name, args.limit,
        )
        if count:
            total_msgs += count
            total_sessions += 1
            if total_sessions % 50 == 0:
                print(f"  {total_sessions} 个会话，{total_msgs:,} 条消息...")

    print(f"\n✅ 完成！会话: {total_sessions}，消息: {total_msgs:,}")
    print(f"📁 输出: {data_root.resolve()}")


if __name__ == "__main__":
    main()
